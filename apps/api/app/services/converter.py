"""
Document conversion engine — comprehensive any-to-any.

Input formats (25+):
  Documents: PDF, DOCX, DOC, PPTX, XLSX, XLS, ODT, ODS, ODP, RTF, EPUB
  Text:      HTML, HTM, TXT, CSV, MD, XML, JSON, YAML, TSV
  Images:    PNG, JPG, JPEG, WebP, BMP, TIFF, TIF, SVG

Output formats (9):
  markdown, pdf, html, txt, docx, csv, json, xml, epub

Strategy: input → Markdown (intermediate) → output renderer
"""

import asyncio
import csv as csv_mod
import io as _io
import json
import os
import re
import tempfile
from html.parser import HTMLParser
from pathlib import Path

import markdown as md
import yaml
from docling.datamodel.base_models import InputFormat
from docling.datamodel.pipeline_options import PdfPipelineOptions, TableFormerMode
from docling.document_converter import DocumentConverter, PdfFormatOption
from docx import Document as DocxDocument
from docx.oxml.ns import qn
from weasyprint import HTML

from app.models import ConversionOptions


# ═══════════════════════════════════════════════════════════════
# Format registry
# ═══════════════════════════════════════════════════════════════

INPUT_FORMATS = {
    # Documents
    ".pdf", ".docx", ".doc", ".pptx", ".xlsx", ".xls",
    ".odt", ".ods", ".odp", ".rtf", ".epub",
    # Text / data
    ".html", ".htm", ".txt", ".csv", ".tsv", ".md",
    ".xml", ".json", ".yaml", ".yml",
    # Images (OCR)
    ".png", ".jpg", ".jpeg", ".webp", ".bmp", ".tiff", ".tif", ".svg",
}

OUTPUT_FORMATS = {"markdown", "pdf", "html", "txt", "docx", "csv", "json", "xml", "epub"}

# CSV output only from tabular sources
CSV_CAPABLE_INPUTS = {".xlsx", ".xls", ".ods", ".csv", ".tsv"}

EXT_TO_MIME = {
    ".pdf": "application/pdf",
    ".docx": "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    ".doc": "application/msword",
    ".pptx": "application/vnd.openxmlformats-officedocument.presentationml.presentation",
    ".xlsx": "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
    ".xls": "application/vnd.ms-excel",
    ".odt": "application/vnd.oasis.opendocument.text",
    ".ods": "application/vnd.oasis.opendocument.spreadsheet",
    ".odp": "application/vnd.oasis.opendocument.presentation",
    ".rtf": "application/rtf",
    ".epub": "application/epub+zip",
    ".html": "text/html", ".htm": "text/html",
    ".txt": "text/plain",
    ".csv": "text/csv", ".tsv": "text/tab-separated-values",
    ".md": "text/markdown",
    ".xml": "application/xml",
    ".json": "application/json",
    ".yaml": "application/x-yaml", ".yml": "application/x-yaml",
    ".png": "image/png", ".jpg": "image/jpeg", ".jpeg": "image/jpeg",
    ".webp": "image/webp", ".bmp": "image/bmp",
    ".tiff": "image/tiff", ".tif": "image/tiff",
    ".svg": "image/svg+xml",
}

OUTPUT_META = {
    "markdown": (".md", "text/markdown"),
    "pdf":      (".pdf", "application/pdf"),
    "html":     (".html", "text/html"),
    "txt":      (".txt", "text/plain"),
    "docx":     (".docx", "application/vnd.openxmlformats-officedocument.wordprocessingml.document"),
    "csv":      (".csv", "text/csv"),
    "json":     (".json", "application/json"),
    "xml":      (".xml", "application/xml"),
    "epub":     (".epub", "application/epub+zip"),
}


def get_supported_conversions() -> dict:
    """Return the full conversion matrix."""
    return {
        "input_formats": sorted(INPUT_FORMATS),
        "output_formats": sorted(OUTPUT_FORMATS),
        "total_input": len(INPUT_FORMATS),
        "total_output": len(OUTPUT_FORMATS),
        "notes": {
            "csv_output": "CSV output only from tabular inputs (XLSX, XLS, ODS, CSV, TSV)",
            "images": "Image inputs use OCR to extract text",
            "strategy": "All conversions route through Markdown as intermediate format",
        },
    }


def can_convert(input_ext: str, output_format: str) -> bool:
    if input_ext not in INPUT_FORMATS:
        return False
    if output_format not in OUTPUT_FORMATS:
        return False
    if output_format == "csv" and input_ext not in CSV_CAPABLE_INPUTS:
        return False
    return True


# ═══════════════════════════════════════════════════════════════
# Main dispatcher
# ═══════════════════════════════════════════════════════════════

async def convert(content: bytes, filename: str, content_type: str,
                  output_format: str, options: ConversionOptions | None = None) -> tuple[bytes, str, str]:
    """Universal converter. Returns (output_bytes, output_filename, output_content_type)."""
    ext = Path(filename).suffix.lower()
    opts = options or ConversionOptions()

    if not can_convert(ext, output_format):
        raise ValueError(f"Cannot convert {ext} to {output_format}")

    out_ext, out_mime = OUTPUT_META[output_format]
    out_filename = Path(filename).stem + out_ext

    if _is_passthrough(ext, output_format):
        return content, out_filename, out_mime

    # Special: tabular → CSV
    if output_format == "csv":
        result = await asyncio.to_thread(_to_csv, content, filename)
        return result.encode("utf-8"), out_filename, out_mime

    # Step 1: Any input → Markdown
    markdown_text = await _to_markdown(content, filename, ext, opts)

    # Step 2: Markdown → target
    if output_format == "markdown":
        return markdown_text.encode("utf-8"), out_filename, out_mime
    elif output_format == "pdf":
        pdf_bytes = await asyncio.to_thread(_markdown_to_pdf, markdown_text, opts)
        return pdf_bytes, out_filename, out_mime
    elif output_format == "html":
        html_str = _markdown_to_html(markdown_text)
        return html_str.encode("utf-8"), out_filename, out_mime
    elif output_format == "txt":
        return _strip_markdown(markdown_text).encode("utf-8"), out_filename, out_mime
    elif output_format == "docx":
        docx_bytes = await asyncio.to_thread(_markdown_to_docx, markdown_text)
        return docx_bytes, out_filename, out_mime
    elif output_format == "json":
        json_str = _markdown_to_json(markdown_text, filename)
        return json_str.encode("utf-8"), out_filename, out_mime
    elif output_format == "xml":
        xml_str = _markdown_to_xml(markdown_text, filename)
        return xml_str.encode("utf-8"), out_filename, out_mime
    elif output_format == "epub":
        epub_bytes = await asyncio.to_thread(_markdown_to_epub, markdown_text, filename)
        return epub_bytes, out_filename, out_mime
    else:
        raise ValueError(f"Unsupported output format: {output_format}")


# Backwards-compatible wrappers
async def convert_to_markdown(content: bytes, filename: str, content_type: str, options: ConversionOptions | None = None) -> str:
    r, _, _ = await convert(content, filename, content_type, "markdown", options)
    return r.decode("utf-8")

async def convert_to_pdf(content: bytes, filename: str, content_type: str, options: ConversionOptions | None = None) -> bytes:
    r, _, _ = await convert(content, filename, content_type, "pdf", options)
    return r

async def convert_to_html(content: bytes, filename: str, content_type: str, options: ConversionOptions | None = None) -> str:
    r, _, _ = await convert(content, filename, content_type, "html", options)
    return r.decode("utf-8")

async def convert_to_txt(content: bytes, filename: str, content_type: str, options: ConversionOptions | None = None) -> str:
    r, _, _ = await convert(content, filename, content_type, "txt", options)
    return r.decode("utf-8")

async def convert_to_docx(content: bytes, filename: str, content_type: str, options: ConversionOptions | None = None) -> bytes:
    r, _, _ = await convert(content, filename, content_type, "docx", options)
    return r


# ═══════════════════════════════════════════════════════════════
# Step 1: Any input → Markdown
# ═══════════════════════════════════════════════════════════════

async def _to_markdown(content: bytes, filename: str, ext: str, opts: ConversionOptions) -> str:
    if ext in (".txt", ".md"):
        return content.decode("utf-8", errors="replace")
    if ext == ".csv":
        return _csv_to_markdown(content, ",")
    if ext == ".tsv":
        return _csv_to_markdown(content, "\t")
    if ext in (".html", ".htm"):
        return await asyncio.to_thread(_html_to_markdown, content)
    if ext == ".docx":
        return await asyncio.to_thread(_docx_to_markdown, content)
    if ext == ".json":
        return _json_to_markdown(content)
    if ext in (".yaml", ".yml"):
        return _yaml_to_markdown(content)
    if ext == ".xml":
        return _xml_to_markdown(content)
    if ext == ".epub":
        return await asyncio.to_thread(_epub_to_markdown, content)
    # Everything else: Docling (PDF, PPTX, XLSX, images, DOC, RTF, ODT, ODS, ODP)
    return await asyncio.to_thread(_docling_convert, content, filename, opts)


# ═══════════════════════════════════════════════════════════════
# Step 2: Markdown → Target renderers
# ═══════════════════════════════════════════════════════════════

def _markdown_to_html(markdown_text: str) -> str:
    body = md.markdown(markdown_text, extensions=["tables", "fenced_code", "codehilite"])
    return f"""<!DOCTYPE html>
<html><head><meta charset="utf-8">
<style>
  body {{ font-family: 'Helvetica Neue', Arial, sans-serif; max-width: 800px; margin: 40px auto; line-height: 1.6; color: #1a1a1a; }}
  h1,h2,h3 {{ color: #0f172a; }} code {{ background: #f1f5f9; padding: 2px 6px; border-radius: 4px; }}
  pre {{ background: #f1f5f9; padding: 16px; border-radius: 8px; overflow-x: auto; }}
  table {{ border-collapse: collapse; width: 100%; }} th,td {{ border: 1px solid #e2e8f0; padding: 8px 12px; }}
  th {{ background: #f8fafc; }} img {{ max-width: 100%; }}
</style>
</head><body>{body}</body></html>"""


def _markdown_to_pdf(markdown_text: str, opts: ConversionOptions) -> bytes:
    margin = opts.pdf_margin or "40px"
    page_size = opts.pdf_page_size or "A4"
    body = md.markdown(markdown_text, extensions=["tables", "fenced_code", "codehilite"])
    styled = f"""<!DOCTYPE html><html><head><meta charset="utf-8">
<style>@page {{ size: {page_size}; margin: {margin}; }}
body {{ font-family: 'Helvetica Neue', Arial, sans-serif; line-height: 1.6; color: #1a1a1a; }}
h1,h2,h3 {{ color: #0f172a; }} code {{ background: #f1f5f9; padding: 2px 6px; border-radius: 4px; }}
pre {{ background: #f1f5f9; padding: 16px; border-radius: 8px; }} table {{ border-collapse: collapse; width: 100%; }}
th,td {{ border: 1px solid #e2e8f0; padding: 8px 12px; }} th {{ background: #f8fafc; }} img {{ max-width: 100%; }}
</style></head><body>{body}</body></html>"""
    return HTML(string=styled).write_pdf()


def _markdown_to_docx(markdown_text: str) -> bytes:
    doc = DocxDocument()
    for line in markdown_text.split("\n"):
        s = line.strip()
        if not s:
            continue
        if s.startswith("#### "):
            doc.add_heading(s[5:], level=4)
        elif s.startswith("### "):
            doc.add_heading(s[4:], level=3)
        elif s.startswith("## "):
            doc.add_heading(s[3:], level=2)
        elif s.startswith("# "):
            doc.add_heading(s[2:], level=1)
        elif s.startswith(("- ", "* ")):
            doc.add_paragraph(s[2:], style="List Bullet")
        elif re.match(r"^\d+\.\s", s):
            doc.add_paragraph(re.sub(r"^\d+\.\s", "", s), style="List Number")
        elif s.startswith("|"):
            cells = [c.strip() for c in s.strip("|").split("|") if c.strip() and not re.match(r"^-+$", c.strip())]
            if cells:
                doc.add_paragraph(" | ".join(cells))
        else:
            doc.add_paragraph(s)
    buf = _io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def _markdown_to_json(markdown_text: str, filename: str) -> str:
    """Convert markdown to structured JSON."""
    sections = []
    current_section = {"heading": "", "level": 0, "content": []}

    for line in markdown_text.split("\n"):
        heading_match = re.match(r"^(#{1,6})\s+(.+)", line)
        if heading_match:
            if current_section["heading"] or current_section["content"]:
                sections.append(current_section)
            current_section = {
                "heading": heading_match.group(2),
                "level": len(heading_match.group(1)),
                "content": [],
            }
        elif line.strip():
            current_section["content"].append(line.strip())

    if current_section["heading"] or current_section["content"]:
        sections.append(current_section)

    result = {
        "source": filename,
        "format": "dockonvert-structured",
        "sections": sections,
        "raw_markdown": markdown_text,
    }
    return json.dumps(result, indent=2, ensure_ascii=False)


def _markdown_to_xml(markdown_text: str, filename: str) -> str:
    """Convert markdown to XML."""
    from xml.sax.saxutils import escape
    lines = ['<?xml version="1.0" encoding="UTF-8"?>', f'<document source="{escape(filename)}">']

    for line in markdown_text.split("\n"):
        heading_match = re.match(r"^(#{1,6})\s+(.+)", line)
        if heading_match:
            level = len(heading_match.group(1))
            lines.append(f'  <heading level="{level}">{escape(heading_match.group(2))}</heading>')
        elif line.startswith(("- ", "* ")):
            lines.append(f"  <list-item>{escape(line[2:])}</list-item>")
        elif line.startswith("|"):
            cells = [c.strip() for c in line.strip("|").split("|") if c.strip() and not re.match(r"^-+$", c.strip())]
            if cells:
                row = "".join(f"<cell>{escape(c)}</cell>" for c in cells)
                lines.append(f"  <table-row>{row}</table-row>")
        elif line.strip():
            lines.append(f"  <paragraph>{escape(line.strip())}</paragraph>")

    lines.append("</document>")
    return "\n".join(lines)


def _markdown_to_epub(markdown_text: str, filename: str) -> bytes:
    """Convert markdown to EPUB."""
    from ebooklib import epub

    book = epub.EpubBook()
    title = Path(filename).stem
    book.set_identifier(f"dockonvert-{title}")
    book.set_title(title)
    book.set_language("en")
    book.add_author("DocKonvert")

    html_body = md.markdown(markdown_text, extensions=["tables", "fenced_code"])
    chapter = epub.EpubHtml(title=title, file_name="content.xhtml", lang="en")
    chapter.content = f"<html><body>{html_body}</body></html>"
    book.add_item(chapter)

    book.toc = [chapter]
    book.add_item(epub.EpubNcx())
    book.add_item(epub.EpubNav())
    book.spine = ["nav", chapter]

    style = epub.EpubItem(uid="style", file_name="style.css", media_type="text/css",
                          content=b"body { font-family: serif; line-height: 1.6; } code { font-family: monospace; }")
    book.add_item(style)
    chapter.add_item(style)

    buf = _io.BytesIO()
    epub.write_epub(buf, book)
    return buf.getvalue()


# ═══════════════════════════════════════════════════════════════
# Input parsers: specific format → Markdown
# ═══════════════════════════════════════════════════════════════

def _docling_convert(content: bytes, filename: str, opts: ConversionOptions) -> str:
    """PDF, PPTX, XLSX, images, DOC, RTF, ODT, ODS, ODP → Markdown via Docling."""
    ext = Path(filename).suffix.lower()
    with tempfile.NamedTemporaryFile(delete=False, suffix=ext) as tmp:
        tmp.write(content)
        tmp_path = tmp.name
    try:
        pdf_options = PdfPipelineOptions()
        pdf_options.do_ocr = opts.ocr if opts.ocr is not None else True
        pdf_options.do_table_structure = True
        pdf_options.table_structure_options.mode = (
            TableFormerMode.FAST if opts.table_mode == "fast" else TableFormerMode.ACCURATE
        )
        pdf_options.generate_picture_images = opts.include_images if opts.include_images is not None else True
        converter = DocumentConverter(
            format_options={InputFormat.PDF: PdfFormatOption(pipeline_options=pdf_options)}
        )
        result = converter.convert(tmp_path)
        return result.document.export_to_markdown()
    finally:
        os.unlink(tmp_path)


def _docx_to_markdown(content: bytes) -> str:
    """DOCX → Markdown with hyperlink preservation."""
    doc = DocxDocument(_io.BytesIO(content))
    lines = []
    for para in doc.paragraphs:
        style = para.style.name.lower() if para.style else ""
        text = _extract_paragraph_text(para)
        if not text.strip():
            lines.append("")
            continue
        if "heading 1" in style:
            lines.append(f"# {text}")
        elif "heading 2" in style:
            lines.append(f"## {text}")
        elif "heading 3" in style:
            lines.append(f"### {text}")
        elif "heading 4" in style:
            lines.append(f"#### {text}")
        elif "list" in style:
            lines.append(f"- {text}")
        else:
            lines.append(text)
    for table in doc.tables:
        lines.append("")
        for i, row in enumerate(table.rows):
            cells = [cell.text.strip() for cell in row.cells]
            lines.append("| " + " | ".join(cells) + " |")
            if i == 0:
                lines.append("| " + " | ".join(["---"] * len(cells)) + " |")
        lines.append("")
    return "\n".join(lines)


def _extract_paragraph_text(para) -> str:
    parts = []
    for child in para._element:
        if child.tag == qn("w:r"):
            for t in child.findall(qn("w:t")):
                parts.append(t.text or "")
        elif child.tag == qn("w:hyperlink"):
            link_text = ""
            for run in child.findall(qn("w:r")):
                for t in run.findall(qn("w:t")):
                    link_text += t.text or ""
            url = ""
            r_id = child.get(qn("r:id"))
            if r_id:
                try:
                    rel = para.part.rels.get(r_id) or para.part.rels[r_id]
                    url = rel.target_ref
                except (KeyError, AttributeError):
                    pass
            if not url:
                anchor = child.get(qn("w:anchor"))
                if anchor:
                    url = f"#{anchor}"
            if url and link_text:
                display = url.replace("mailto:", "") if url.startswith("mailto:") else url
                if link_text == display or not url.startswith(("http", "mailto")):
                    parts.append(link_text)
                else:
                    parts.append(f"[{link_text}]({url})")
            elif link_text:
                parts.append(link_text)
            elif url:
                parts.append(url.replace("mailto:", "") if url.startswith("mailto:") else url)
    return "".join(parts)


def _html_to_markdown(content: bytes) -> str:
    html_text = content.decode("utf-8", errors="replace")
    lines = []
    current_tag = ""

    class _P(HTMLParser):
        def handle_starttag(self, tag, attrs):
            nonlocal current_tag
            current_tag = tag

        def handle_data(self, data):
            text = data.strip()
            if not text:
                return
            tag_map = {"h1": "# ", "h2": "## ", "h3": "### ", "h4": "#### ", "li": "- "}
            prefix = tag_map.get(current_tag, "")
            if current_tag in ("code", "pre"):
                lines.append(f"`{text}`")
            else:
                lines.append(f"{prefix}{text}")

    _P().feed(html_text)
    return "\n".join(lines)


def _csv_to_markdown(content: bytes, delimiter: str = ",") -> str:
    text = content.decode("utf-8", errors="replace")
    reader = csv_mod.reader(text.splitlines(), delimiter=delimiter)
    lines = []
    for i, row in enumerate(reader):
        lines.append("| " + " | ".join(row) + " |")
        if i == 0:
            lines.append("| " + " | ".join(["---"] * len(row)) + " |")
    return "\n".join(lines)


def _json_to_markdown(content: bytes) -> str:
    text = content.decode("utf-8", errors="replace")
    try:
        data = json.loads(text)
    except json.JSONDecodeError:
        return f"```json\n{text}\n```"

    if isinstance(data, list) and data and isinstance(data[0], dict):
        # Array of objects → markdown table
        headers = list(data[0].keys())
        lines = ["| " + " | ".join(headers) + " |"]
        lines.append("| " + " | ".join(["---"] * len(headers)) + " |")
        for row in data:
            lines.append("| " + " | ".join(str(row.get(h, "")) for h in headers) + " |")
        return "\n".join(lines)

    return f"```json\n{json.dumps(data, indent=2, ensure_ascii=False)}\n```"


def _yaml_to_markdown(content: bytes) -> str:
    text = content.decode("utf-8", errors="replace")
    try:
        data = yaml.safe_load(text)
    except yaml.YAMLError:
        return f"```yaml\n{text}\n```"

    if isinstance(data, dict):
        lines = []
        for key, value in data.items():
            if isinstance(value, (dict, list)):
                lines.append(f"## {key}")
                lines.append(f"```yaml\n{yaml.dump(value, default_flow_style=False).strip()}\n```")
            else:
                lines.append(f"**{key}**: {value}")
        return "\n\n".join(lines)

    return f"```yaml\n{text}\n```"


def _xml_to_markdown(content: bytes) -> str:
    import xml.etree.ElementTree as ET
    text = content.decode("utf-8", errors="replace")
    try:
        root = ET.fromstring(text)
    except ET.ParseError:
        return f"```xml\n{text}\n```"

    lines = [f"# {root.tag}"]

    def _walk(elem, depth=0):
        prefix = "#" * min(depth + 2, 6) + " " if depth < 5 else "- "
        if elem.text and elem.text.strip():
            lines.append(f"{prefix}{elem.tag}: {elem.text.strip()}")
        elif len(elem):
            lines.append(f"{prefix}{elem.tag}")
        for child in elem:
            _walk(child, depth + 1)

    for child in root:
        _walk(child)
    return "\n".join(lines)


def _epub_to_markdown(content: bytes) -> str:
    """EPUB → Markdown by extracting HTML chapters."""
    from ebooklib import epub

    book = epub.read_epub(_io.BytesIO(content))
    parts = []

    for item in book.get_items_of_type(9):  # ITEM_DOCUMENT
        html_content = item.get_content()
        md_text = _html_to_markdown(html_content)
        if md_text.strip():
            parts.append(md_text)

    return "\n\n---\n\n".join(parts) if parts else "(Empty EPUB)"


# ═══════════════════════════════════════════════════════════════
# Tabular → CSV
# ═══════════════════════════════════════════════════════════════

def _to_csv(content: bytes, filename: str) -> str:
    ext = Path(filename).suffix.lower()
    if ext == ".csv":
        return content.decode("utf-8", errors="replace")
    if ext == ".tsv":
        text = content.decode("utf-8", errors="replace")
        return text.replace("\t", ",")
    # XLSX/XLS/ODS
    try:
        import openpyxl
        wb = openpyxl.load_workbook(_io.BytesIO(content), read_only=True)
        ws = wb.active
        lines = []
        for row in ws.iter_rows(values_only=True):
            lines.append(",".join(str(c) if c is not None else "" for c in row))
        return "\n".join(lines)
    except Exception:
        opts = ConversionOptions()
        md_text = _docling_convert(content, filename, opts)
        return _strip_markdown(md_text)


# ═══════════════════════════════════════════════════════════════
# Utilities
# ═══════════════════════════════════════════════════════════════

def _is_passthrough(ext: str, output_format: str) -> bool:
    mapping = {".md": "markdown", ".pdf": "pdf", ".html": "html", ".htm": "html",
               ".txt": "txt", ".docx": "docx", ".csv": "csv", ".json": "json",
               ".xml": "xml", ".epub": "epub"}
    return mapping.get(ext) == output_format


def _strip_markdown(text: str) -> str:
    text = re.sub(r"#{1,6}\s+", "", text)
    text = re.sub(r"\*\*(.+?)\*\*", r"\1", text)
    text = re.sub(r"\*(.+?)\*", r"\1", text)
    text = re.sub(r"\[(.+?)\]\(.+?\)", r"\1", text)
    text = re.sub(r"!\[.*?\]\(.+?\)", "", text)
    text = re.sub(r"`(.+?)`", r"\1", text)
    text = re.sub(r"^[-*]\s+", "", text, flags=re.MULTILINE)
    text = re.sub(r"^\d+\.\s+", "", text, flags=re.MULTILINE)
    text = re.sub(r"\|", " ", text)
    text = re.sub(r"---+", "", text)
    return text.strip()
